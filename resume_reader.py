from docx import Document
import time
import tkinter as tk
from tkinter import filedialog

def read_docx(file_path):
    document = Document(file_path)
    full_text = []
    for para in document.paragraphs:
        full_text.append(para.text)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                full_text.append(cell.text)

    return '\n'.join(full_text)


def get_file():
    root = tk.Tk()
    root.withdraw()  # Hide the root window

    file_path = filedialog.askopenfilename(
        title="Select a .docx file",
        filetypes=[("Word files", "*.docx")]
    )

    return file_path if file_path else None



# parse important information from the resume
def parse_resume(content):
    _experience = []
    _education = []
    _skills = []

    lines = content.split('\n')
    for line in lines:
        if "Experience".upper() in line:

            for exp in lines[lines.index(line)+1:]:
                if exp.strip() == "Education".upper() or exp.strip() == "Skills".upper():
                    # make sure to break out of the loop if we reach the next section
                    break

                if exp.strip() == "":
                    break
                # save the experience in a list
                _experience.append(exp.strip())

        if "Education".upper() in line:
            for edu in lines[lines.index(line)+1:]:
                if edu.strip() == "Experience".upper() or edu.strip() == "Skills".upper():
                    # make sure to break out of the loop if we reach the next section
                    break

                if edu.strip() == "":
                    break
                # save the education in a list
                _education.append(edu.strip())

        if "Skills".upper() in line:
            for skill in lines[lines.index(line)+1:]:
                if skill.strip() == "Experience".upper() or skill.strip() == "Education".upper():
                    # make sure to break out of the loop if we reach the next section
                    break

                if skill.strip() == "":
                    break
                # save the skills in a list
                _skills.append(skill.strip())

    print("Experience:")
    for exp in _experience:
        print(f"- {exp}")

    print("\nEducation:")
    for edu in _education:
        print(f"- {edu}")

    print("\nSkills:")
    for skill in _skills:
        print(f"- {skill}")

def main():
    print("Welcome to the Resume Reader!")
    print("Please select a .docx file to read.")
    print()

    time.sleep(1)

    file_path = get_file()
    if file_path:
        content = read_docx(file_path)
        parse_resume(content)
    else:
        print("No file selected.")

if __name__ == "__main__":
    main()
